#!/usr/bin/env python3
"""Export a translated PDF as an editable, block-addressable Word companion."""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Any

from common import SCHEMA_VERSION, load_fitz, normalize_space, require_file, sha256_file, stable_id, utc_now, write_json
from toc_common import analyze_page, detection_score


def load_docx():
    try:
        import docx
        from docx.enum.section import WD_SECTION
        from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise SystemExit(
            "python-docx is required. Install it in an isolated environment with: "
            "python -m pip install python-docx"
        ) from exc
    return {
        "docx": docx, "WD_SECTION": WD_SECTION,
        "WD_ALIGN_VERTICAL": WD_ALIGN_VERTICAL,
        "WD_TABLE_ALIGNMENT": WD_TABLE_ALIGNMENT,
        "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
        "OxmlElement": OxmlElement, "qn": qn,
        "Inches": Inches, "Pt": Pt, "RGBColor": RGBColor,
    }


def set_cell_margins(cell: Any, api: dict[str, Any], top: int = 80, start: int = 120,
                     bottom: int = 80, end: int = 120) -> None:
    OxmlElement, qn = api["OxmlElement"], api["qn"]
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table: Any, widths: list[int], api: dict[str, Any]) -> None:
    OxmlElement, qn = api["OxmlElement"], api["qn"]
    total = sum(widths)
    table.autofit = False
    table.alignment = api["WD_TABLE_ALIGNMENT"].LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell, api)
            cell.vertical_alignment = api["WD_ALIGN_VERTICAL"].CENTER


def repeat_header(row: Any, api: dict[str, Any]) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = api["OxmlElement"]("w:tblHeader")
    tbl_header.set(api["qn"]("w:val"), "true")
    tr_pr.append(tbl_header)


def shade_cell(cell: Any, fill: str, api: dict[str, Any]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(api["qn"]("w:shd"))
    if shd is None:
        shd = api["OxmlElement"]("w:shd")
        tc_pr.append(shd)
    shd.set(api["qn"]("w:fill"), fill)


def set_run_font(run: Any, api: dict[str, Any], name: str = "Calibri", size: float = 11,
                 bold: bool | None = None, color: str | None = None,
                 east_asia_name: str = "Arial Unicode MS") -> None:
    effective_name = east_asia_name if not run.text.isascii() else name
    run.font.name = effective_name
    run._element.get_or_add_rPr().rFonts.set(api["qn"]("w:ascii"), effective_name)
    run._element.get_or_add_rPr().rFonts.set(api["qn"]("w:hAnsi"), effective_name)
    run._element.get_or_add_rPr().rFonts.set(api["qn"]("w:eastAsia"), effective_name)
    run.font.size = api["Pt"](size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = api["RGBColor"].from_string(color)


def add_hidden_id(cell: Any, item_id: str, api: dict[str, Any]) -> None:
    paragraph = cell.paragraphs[0]
    hidden = paragraph.add_run(f" [[ID:{item_id}]]")
    hidden.font.hidden = True
    set_run_font(hidden, api, size=6, color="7A7A7A")


def set_cell_text(cell: Any, text: str, api: dict[str, Any], *, bold: bool = False,
                  color: str | None = None, size: float = 10.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = api["Pt"](0)
    paragraph.paragraph_format.space_after = api["Pt"](2)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run, api, size=size, bold=bold, color=color)


def extract_blocks(page: Any, page_number: int) -> list[dict[str, Any]]:
    blocks = []
    for block in page.get_text("blocks", sort=True):
        if len(block) < 7 or int(block[6]) != 0:
            continue
        text = normalize_space(str(block[4]))
        if not text:
            continue
        bbox = [round(float(value), 3) for value in block[:4]]
        block_id = stable_id("block", page_number, *bbox, text, length=14)
        blocks.append({"block_id": block_id, "page": page_number, "bbox": bbox, "text": text})
    return blocks


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("translated_pdf")
    parser.add_argument("--output", required=True)
    parser.add_argument("--manifest-out", required=True)
    parser.add_argument("--title", default="Editable Translation Companion")
    parser.add_argument("--language", default="zh-CN")
    args = parser.parse_args()

    api = load_docx()
    fitz = load_fitz()
    pdf_path = require_file(args.translated_pdf, "translated PDF")
    output = Path(args.output).expanduser().resolve()
    if output.exists():
        raise SystemExit(f"Output already exists: {output}")
    document_pdf = fitz.open(pdf_path)
    doc = api["docx"].Document()
    section = doc.sections[0]
    section.page_width = api["Inches"](8.5)
    section.page_height = api["Inches"](11)
    section.top_margin = api["Inches"](1)
    section.right_margin = api["Inches"](1)
    section.bottom_margin = api["Inches"](1)
    section.left_margin = api["Inches"](1)
    section.header_distance = api["Inches"](0.492)
    section.footer_distance = api["Inches"](0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(api["qn"]("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(api["qn"]("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(api["qn"]("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = api["Pt"](11)
    normal.paragraph_format.space_after = api["Pt"](6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, "2E74B5"),
        ("Heading 2", 13, 14, 7, "2E74B5"),
        ("Heading 3", 12, 10, 5, "1F4D78"),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(api["qn"]("w:eastAsia"), "Arial Unicode MS")
        style.font.size = api["Pt"](size)
        style.font.color.rgb = api["RGBColor"].from_string(color)
        style.paragraph_format.space_before = api["Pt"](before)
        style.paragraph_format.space_after = api["Pt"](after)

    header = section.header.paragraphs[0]
    header.alignment = api["WD_ALIGN_PARAGRAPH"].LEFT
    set_run_font(header.add_run("Editable PDF Translation"), api, size=9, color="6B7280")
    footer = section.footer.paragraphs[0]
    footer.alignment = api["WD_ALIGN_PARAGRAPH"].RIGHT
    set_run_font(footer.add_run("PDF remains the layout-authoritative version"), api, size=8, color="6B7280")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = api["Pt"](4)
    set_run_font(title.add_run(args.title), api, size=23, bold=True, color="0B2545")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = api["Pt"](14)
    set_run_font(
        subtitle.add_run(f"{pdf_path.name} | {len(document_pdf)} PDF pages | {args.language}"),
        api, size=10.5, color="555555",
    )
    note = doc.add_paragraph()
    note.paragraph_format.space_after = api["Pt"](14)
    set_run_font(
        note.add_run(
            "Use this Word file for text corrections. Keep hidden block IDs intact so edits "
            "can be exported and reconciled with the layout-controlled PDF."
        ),
        api, size=10.5, color="1F3A5F",
    )

    manifest_items = []
    for page_index, page in enumerate(document_pdf):
        page_number = page_index + 1
        heading = doc.add_paragraph(style="Heading 1")
        heading.paragraph_format.keep_with_next = True
        heading.add_run(f"PDF page {page_number}")
        toc_analysis = analyze_page(page, page_number)
        if detection_score(toc_analysis) >= 0.65 and len(toc_analysis["rows"]) >= 4:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            headers = ("Section", "Editable title", "Destination")
            for index, value in enumerate(headers):
                set_cell_text(table.rows[0].cells[index], value, api, bold=True, color="0B2545")
                shade_cell(table.rows[0].cells[index], "E8EEF5", api)
            repeat_header(table.rows[0], api)
            for toc_row in toc_analysis["rows"]:
                cells = table.add_row().cells
                set_cell_text(cells[0], toc_row["section_number"], api, size=9.5)
                add_hidden_id(cells[0], toc_row["row_id"], api)
                set_cell_text(cells[1], toc_row["source_title"], api)
                set_cell_text(cells[2], toc_row["destination_label"], api, size=9.5)
                cells[2].paragraphs[0].alignment = api["WD_ALIGN_PARAGRAPH"].CENTER
                manifest_items.append({
                    "item_id": toc_row["row_id"], "page": page_number,
                    "kind": "toc-row", "text": toc_row["source_title"],
                })
            set_table_geometry(table, [1200, 6960, 1200], api)
        else:
            blocks = extract_blocks(page, page_number)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            for index, value in enumerate(("Block ID", "Editable translated text")):
                set_cell_text(table.rows[0].cells[index], value, api, bold=True, color="0B2545")
                shade_cell(table.rows[0].cells[index], "E8EEF5", api)
            repeat_header(table.rows[0], api)
            for block in blocks:
                cells = table.add_row().cells
                visible_id = block["block_id"].split("-", 1)[-1][:8]
                set_cell_text(cells[0], visible_id, api, size=8.5, color="555555")
                add_hidden_id(cells[0], block["block_id"], api)
                set_cell_text(cells[1], block["text"], api)
                manifest_items.append({
                    "item_id": block["block_id"], "page": page_number,
                    "kind": "text-block", "bbox": block["bbox"], "text": block["text"],
                })
            set_table_geometry(table, [1600, 7760], api)
        if page_number < len(document_pdf):
            doc.add_page_break()

    doc.core_properties.title = args.title
    doc.core_properties.subject = "Editable companion for a layout-controlled translated PDF"
    doc.core_properties.keywords = "PDF translation, editable companion, block IDs"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    document_pdf.close()
    manifest = {
        "schema_version": SCHEMA_VERSION,
        "tool": "export_editable_docx",
        "generated_at": utc_now(),
        "source": {"path": str(pdf_path), "sha256": sha256_file(pdf_path)},
        "output": {"path": str(output), "sha256": sha256_file(output)},
        "style_preset": "compact_reference_guide",
        "header_pattern": "memo_masthead",
        "items": manifest_items,
    }
    write_json(Path(args.manifest_out).expanduser().resolve(), manifest)
    print(f"Wrote editable DOCX with {len(manifest_items)} addressable items")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
